"""
Resume Parser Service for extracting text from various file formats
"""
import logging
import os
import hashlib
import tempfile
from typing import BinaryIO, Tuple, Optional
from uuid import UUID
import mimetypes
from app.utils.document_utils import extract_text_from_bytes, get_mime_type_from_filename

from bs4 import BeautifulSoup

logger = logging.getLogger(__name__)


class ResumeParserService:
    """Service for parsing resume files and extracting text content"""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.txt', '.doc', '.docx'}
    SUPPORTED_MIME_TYPES = {
        'application/pdf',
        'text/plain',
        'application/msword',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    }
    
    def __init__(self):
        self.max_file_size = 10 * 1024 * 1024  # 10MB
    
    def validate_file(self, filename: str, file_size: int) -> Tuple[bool, str]:
        """验证文件是否符合要求"""
        # 检查文件大小
        if file_size > self.max_file_size:
            return False, f"文件大小超过限制 ({self.max_file_size / 1024 / 1024}MB)"
        
        # 检查文件扩展名
        _, ext = os.path.splitext(filename.lower())
        if ext not in self.SUPPORTED_EXTENSIONS:
            return False, f"不支持的文件格式。支持的格式: {', '.join(self.SUPPORTED_EXTENSIONS)}"
        
        return True, "文件验证通过"
    
    async def extract_text_from_file(self, file_content: bytes, filename: str) -> str:
        """统一委派到文档工具模块进行文本抽取"""
        try:
            return extract_text_from_bytes(file_content, filename)
        except Exception as e:
            logger.error(f"提取文本内容失败: {e}")
            raise Exception(f"文件解析失败: {str(e)}")
    
    async def _extract_from_txt(self, file_content: bytes) -> str:
        """从TXT文件提取文本"""
        try:
            # 尝试不同的编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
            
            for encoding in encodings:
                try:
                    return file_content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            
            # 如果所有编码都失败，使用utf-8并忽略错误
            return file_content.decode('utf-8', errors='ignore')
            
        except Exception as e:
            logger.error(f"TXT文件解析失败: {e}")
            raise
    
    async def _extract_from_pdf(self, file_content: bytes) -> str:
        """从PDF文件提取文本"""
        try:
            # 方法1: 使用PyPDF2
            try:
                import PyPDF2
                import io
                
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                text_content = []
                
                for page in pdf_reader.pages:
                    try:
                        text = page.extract_text()
                        if text and text.strip():
                            text_content.append(text.strip())
                    except Exception as e:
                        logger.warning(f"PDF页面提取失败: {e}")
                        continue
                
                result = '\n'.join(text_content)
                if result.strip():
                    return result
                    
            except Exception as e:
                logger.warning(f"PyPDF2解析失败，尝试pdfplumber: {e}")
            
            # 方法2: 使用pdfplumber
            try:
                import pdfplumber
                import io
                
                with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                    text_content = []
                    for page_num, page in enumerate(pdf.pages):
                        try:
                            text = page.extract_text()
                            if text and text.strip():
                                text_content.append(text.strip())
                        except Exception as e:
                            logger.warning(f"PDF页面 {page_num + 1} 提取失败: {e}")
                            continue
                    
                    result = '\n'.join(text_content)
                    if result.strip():
                        return result
                        
            except Exception as e:
                logger.warning(f"pdfplumber解析失败: {e}")
            
            # 如果以上方法都失败，返回提示
            return ""
            
        except ImportError:
            logger.error("PDF解析库未安装，请安装 PyPDF2 或 pdfplumber")
            raise Exception("PDF解析功能不可用，请联系管理员")
            
        except Exception as e:
            logger.error(f"PDF文件解析失败: {e}")
            raise Exception(f"PDF文件解析失败: {str(e)}")

    async def extract_text_with_advanced_formatting(self, html_content: str) -> str:
        """高级格式保留：区分段落、列表、标题等"""
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(html_content, 'html.parser')

        # 移除不需要的元素
        for element in soup(["script", "style", "head", "title", "meta"]):
            element.decompose()

        # 处理特定标签的格式
        formatting_rules = {
            'p': '\n\n',  # 段落：双换行
            'br': '\n',  # 换行：单换行
            'div': '\n',  # 分区：单换行
            'li': '\n• ',  # 列表项：换行+项目符号
            'h1': '\n\n# ',  # 一级标题
            'h2': '\n\n## ',  # 二级标题
            'h3': '\n\n### ',  # 三级标题
            'blockquote': '\n> '  # 引用块
        }

        # 应用格式规则
        for tag, prefix in formatting_rules.items():
            for element in soup.find_all(tag):
                # 在元素内容前添加格式前缀
                element.insert_before(prefix)

        # 获取文本
        text = soup.get_text()

        # 清理文本：合并多余空行，保留合理的段落间距
        lines = []
        empty_line_count = 0
        for line in text.splitlines():
            stripped_line = line.strip()
            if stripped_line:
                lines.append(stripped_line)
                empty_line_count = 0
            elif empty_line_count < 2:  # 最多保留两个连续空行
                lines.append('')
                empty_line_count += 1

        return '\n'.join(lines)

    async def _extract_from_doc(self, file_content: bytes) -> str:
        """从DOC文件提取文本（简化：仅使用 docx2txt）"""
        try:

            with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name
            with open(temp_file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()

            text = await self.extract_text_with_advanced_formatting(html_content)
            return text
            # 提取纯文本


        except Exception as e:
            logger.error(f"DOC文件解析失败: {e}")
            raise Exception(f"DOC文件解析失败: {str(e)}")
    
    async def _extract_from_docx(self, file_content: bytes) -> str:
        """从DOCX文件提取文本（简化：仅使用 python-docx）"""
        try:
            from docx import Document
            import io

            doc = Document(io.BytesIO(file_content))
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)
            result = '\n'.join(text_content)
            return result.strip() if result.strip() else ""
        except ImportError:
            logger.error("DOCX解析库未安装，请安装 python-docx")
            raise Exception("DOCX解析功能不可用，请联系管理员")
        except Exception as e:
            logger.error(f"DOCX文件解析失败: {e}")
            raise Exception(f"DOCX文件解析失败: {str(e)}")
    
    def _detect_document_type(self, file_content: bytes, filename: str) -> str:
        """检测文档类型和特征"""
        try:
            ext = filename.lower().split('.')[-1]
            
            if ext == 'pdf':
                # 检测PDF是否可能是扫描文档
                try:
                    import PyPDF2
                    import io
                    
                    pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                    total_pages = len(pdf_reader.pages)
                    text_pages = 0
                    
                    for page in pdf_reader.pages:
                        try:
                            text = page.extract_text()
                            if text and text.strip():
                                text_pages += 1
                        except:
                            continue
                    
                    # 如果少于30%的页面有文本，可能是扫描文档
                    if total_pages > 0 and text_pages / total_pages < 0.3:
                        return "scanned_pdf"
                        
                except:
                    pass
                    
            elif ext in ['doc', 'docx']:
                # 检测Word文档是否可能是图片文档
                try:
                    import zipfile
                    import io
                    
                    with zipfile.ZipFile(io.BytesIO(file_content)) as docx:
                        # 检查是否包含大量图片
                        image_files = [f for f in docx.namelist() if f.startswith('word/media/')]
                        if len(image_files) > 5:  # 如果有很多图片，可能是图文混排
                            return "image_heavy_doc"
                            
                except:
                    pass
                    
            return "normal"
            
        except Exception as e:
            logger.warning(f"文档类型检测失败: {e}")
            return "unknown"

    def _validate_extracted_content(self, text: str, filename: str) -> dict:
        """验证提取的文本内容质量"""
        if not text or not text.strip():
            return {
                'is_valid': False,
                'reason': 'empty_content',
                'message': '未提取到任何文本内容'
            }
        
        text = text.strip()
        
        # 检查文本长度
        if len(text) < 10:
            return {
                'is_valid': False,
                'reason': 'too_short',
                'message': f'提取的文本内容太短（{len(text)}字符），可能不是有效的简历内容'
            }
        
        # 检查是否包含中文字符或英文字符
        has_chinese = any('\u4e00' <= char <= '\u9fff' for char in text)
        has_english = any('a' <= char.lower() <= 'z' for char in text)
        
        if not has_chinese and not has_english:
            return {
                'is_valid': False,
                'reason': 'no_readable_text',
                'message': '提取的内容不包含可读的中英文文本，可能是图片或扫描文档'
            }
        
        # 检查是否包含常见的简历关键词
        resume_keywords = ['姓名', 'name', '电话', 'phone', '邮箱', 'email', '工作', 'work', '经验', 'experience', '教育', 'education', '技能', 'skills']
        text_lower = text.lower()
        keyword_count = sum(1 for keyword in resume_keywords if keyword in text_lower)
        
        if keyword_count < 2:
            logger.warning(f"文件 '{filename}' 提取的内容可能不是简历，缺少常见关键词。找到关键词数量: {keyword_count}")
        
        return {
            'is_valid': True,
            'reason': 'valid_content',
            'message': '内容验证通过',
            'text_length': len(text),
            'has_chinese': has_chinese,
            'has_english': has_english,
            'keyword_count': keyword_count
        }

    def get_file_info(self, filename: str, file_content: bytes) -> dict:
        """获取文件基本信息"""
        file_size = len(file_content)
        file_hash = hashlib.sha256(file_content).hexdigest()
        _, ext = os.path.splitext(filename.lower())
        
        return {
            'filename': filename,
            'file_type': ext.lstrip('.'),
            'file_size': file_size,
            'file_hash': file_hash,
            'mime_type': mimetypes.guess_type(filename)[0]
        }
