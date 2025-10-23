"""
Enhanced document service with LangChain integration for document processing and vectorization
"""
import logging
import os
import hashlib
from typing import List, Optional, Dict, Any, BinaryIO
from uuid import UUID
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, update, delete, desc, func, text
from sqlalchemy.orm import selectinload

# LangChain imports
from langchain_core.documents import Document as LangChainDocument
from langchain_postgres import PGVector

# Document processing imports
import PyPDF2
from docx import Document as DocxDocument
import tempfile

from app.models.document import Document
from app.models.knowledge_base import KnowledgeBase
from app.services.llm_service import LLMService
from app.services.embedding_service import get_embedding_service
from app.schemas.document import DocumentCreate, DocumentUpdate
from app.core.config import settings

logger = logging.getLogger(__name__)


class EnhancedDocumentService:
    """Enhanced service for managing documents with LangChain integration"""

    def __init__(self, db: AsyncSession):
        self.db = db
        self.llm_service = LLMService()

        # Use shared embedding service to avoid repeated initialization
        self.embedding_service = get_embedding_service()
        self.embeddings = self.embedding_service.get_embeddings()
        self.text_splitter = self.embedding_service.get_text_splitter()

        # Database connection string for PGVector (use psycopg2 for sync connection)
        self.connection_string = settings.DATABASE_URL

        logger.info("Enhanced document service initialized with shared embedding service")

    def _split_by_semantic_points(self, text: str, split_points: List[str]) -> List[str]:
        """根据语义分割点切分文本"""
        chunks = []
        current_pos = 0

        # 按顺序查找每个分割点并切分文本
        for point in split_points:
            pos = text.find(point, current_pos)
            if pos != -1:
                # 添加当前位置到分割点位置的文本块
                if pos > current_pos:
                    chunk = text[current_pos:pos].strip()
                    if chunk:
                        chunks.append(chunk)
                current_pos = pos

        # 添加最后一个文本块
        if current_pos < len(text):
            chunk = text[current_pos:].strip()
            if chunk:
                chunks.append(chunk)

        return chunks
    async def get_semantic_split_points(self, content: str) -> List[str]:
        try:
            # 通过聊天模型生成分割点，仅返回以 `~~` 分隔的分割点字符串
            system_prompt = (
                "你是一个文档结构分析助手。只输出用于 split 的分割点字符串，"
                "用`~~`分隔，不要输出任何其他文字。确保每个分割点在原文中唯一，"
                "如果遇到重复标题或目录项，需要在分割点后追加少量后续字符形成唯一片段。"
            )
            user_prompt = (
                "# 任务\n请分析文档，识别适合作为分割点的文本片段。\n\n"
                "# 规则\n"
                "1) 分割点应位于句子或段落的开头；\n"
                "2) 分割后每段尽量<=500字，严禁>1000字；\n"
                "3) 若存在重复片段（例如目录与正文相同标题），需在分割点后追加少量后续内容以确保唯一；\n"
                "4) 仅输出分割点字符串，使用`~~`分隔，不要解释或添加其他文本。\n\n"
                f"# 文档（截断）\n{content[:10000]}"
            )

            response = await self.llm_service.client.chat.completions.create(
                model=self.llm_service.llm_model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                temperature=0.2,
                max_tokens=1000,
            )
            raw = response.choices[0].message.content or ""
            points = [p.strip() for p in raw.split("~~") if p.strip()]
            if not points:
                return []

            # 去重并保持顺序
            seen = set()
            unique_points = []
            for p in points:
                if p not in seen:
                    seen.add(p)
                    unique_points.append(p)

            # 保证每个分割点在正文中唯一：必要时追加后续字符
            def ensure_unique(point: str) -> str:
                start = content.find(point)
                if start == -1:
                    return ""  # 模型输出不在原文中，丢弃
                # 统计出现次数
                count = 0
                search_pos = 0
                while True:
                    idx = content.find(point, search_pos)
                    if idx == -1:
                        break
                    count += 1
                    search_pos = idx + 1
                if count <= 1:
                    return point
                # 重复：逐步扩展片段直到唯一或达到限制
                # 最多追加 100 个字符，步长 10
                max_extra = 100
                step = 10
                extra = 0
                while extra <= max_extra:
                    candidate = content[start:start + len(point) + extra]
                    # 重新统计
                    c = 0
                    sp = 0
                    while True:
                        j = content.find(candidate, sp)
                        if j == -1:
                            break
                        c += 1
                        sp = j + 1
                    if c <= 1:
                        return candidate
                    extra += step
                # 仍不唯一则返回原始（极少数情况），后续切分时按位置处理
                return point

            adjusted_points_with_index: List[tuple[int, str]] = []
            for p in unique_points:
                adj = ensure_unique(p)
                if not adj:
                    continue
                idx = content.find(adj)
                if idx != -1:
                    adjusted_points_with_index.append((idx, adj))

            # 按在正文中的出现位置排序
            adjusted_points_with_index.sort(key=lambda x: x[0])
            final_points = [pt for _, pt in adjusted_points_with_index]
            return final_points
        except Exception as e:
            logger.error(f"Enhanced document service get_semantic_split_points error: {e}")
            return []

    def _force_split_long_chunk(self, chunk: str) -> List[str]:
        """强制分割超长段落（超过1000字符）"""
        max_length = 1000
        chunks = []

        # 先尝试按换行符分割
        if '\n' in chunk:
            lines = chunk.split('\n')
            current_chunk = ""
            for line in lines:
                if len(current_chunk) + len(line) + 1 > max_length:
                    if current_chunk:
                        chunks.append(current_chunk)
                        current_chunk = line
                    else:
                        chunks.append(line[:max_length])
                        current_chunk = line[max_length:]
                else:
                    if current_chunk:
                        current_chunk += "\n" + line
                    else:
                        current_chunk = line
            if current_chunk:
                chunks.append(current_chunk)
        else:
            # 没有换行符则直接按长度分割
            chunks = [chunk[i:i + max_length] for i in range(0, len(chunk), max_length)]

        return chunks

    def _merge_short_chunks(self, chunks: List[str], min_length: int = 50, max_length: int = 1000) -> List[str]:
        """合并过短片段，避免产生极短段落，同时不超过最大长度"""
        if not chunks:
            return []
        merged: List[str] = []
        i = 0
        while i < len(chunks):
            cur = chunks[i]
            if len(cur) < min_length and i + 1 < len(chunks):
                nxt = chunks[i + 1]
                if len(cur) + len(nxt) <= max_length:
                    merged.append(cur + ("\n" if cur and nxt and not cur.endswith("\n") else "") + nxt)
                    i += 2
                    continue
            merged.append(cur)
            i += 1
        return merged

    async def _split_text(self, content: str) -> List[str]:
        """主分割流程：使用 LLM 分割点 + 切分 + 长度约束"""
        # 1) 获取分割点（可能为空）
        points = await self.get_semantic_split_points(content)
        # 2) 根据分割点切分；若为空则整体作为一个片段
        if points:
            chunks = self._split_by_semantic_points(content, points)
        else:
            chunks = [content]
        # 3) 对超长片段进行强制分割（>1000）
        normalized: List[str] = []
        for ch in chunks:
            if len(ch) > 1000:
                normalized.extend(self._force_split_long_chunk(ch))
            else:
                normalized.append(ch)
        # 4) 合并过短片段（<50）
        normalized = self._merge_short_chunks(normalized, min_length=50, max_length=1000)
        # 5) 去除空白
        normalized = [c.strip() for c in normalized if c and c.strip()]
        return normalized

    async def upload_document(
        self,
        file,  # UploadFile object
        user_id: UUID,
        category: Optional[str] = None,
        tags: Optional[List[str]] = None,
        knowledge_base_id: Optional[str] = None
    ) -> Document:
        """Upload and process a document"""
        # Convert knowledge_base_id from string to UUID if provided
        kb_id = None
        if knowledge_base_id:
            try:
                kb_id = UUID(knowledge_base_id)
            except (ValueError, TypeError):
                logger.warning(f"Invalid knowledge_base_id format: {knowledge_base_id}")
                kb_id = None

        return await self.upload_and_process_document(
            file=file.file,
            filename=file.filename,
            user_id=user_id,
            knowledge_base_id=kb_id,
            category=category,
            tags=tags
        )

    async def upload_and_process_document(
        self,
        file: BinaryIO,
        filename: str,
        user_id: UUID,
        knowledge_base_id: Optional[UUID] = None,
        category: Optional[str] = None,
        tags: Optional[List[str]] = None
    ) -> Document:
        """Upload and process document with full text extraction and vectorization"""
        try:
            # Read file content
            file_content = file.read()
            file.seek(0)  # Reset file pointer

            # Generate file hash
            file_hash = hashlib.sha256(file_content).hexdigest()

            # Check if document already exists
            existing_doc = await self._get_document_by_hash(file_hash, user_id)
            if existing_doc:
                logger.info(f"Document with hash {file_hash} already exists")
                return existing_doc

            # Determine MIME type
            mime_type = self._get_mime_type(filename)

            # Save file temporarily for processing
            temp_file_path = await self._save_temp_file(file_content, filename)

            try:
                # Extract text content
                extracted_content = await self._extract_text_content(temp_file_path, mime_type)

                # Generate summary
                summary = await self.llm_service.summarize_text(extracted_content) if extracted_content else None

                # Generate embedding for the document
                # document_embedding = await self.embeddings.aembed_query(extracted_content) if extracted_content else None

                # Save file permanently
                file_path = await self._save_file(file_content, filename, user_id)

                # Create document record
                document = Document(
                    filename=filename,
                    original_filename=filename,
                    file_path=file_path,
                    file_size=len(file_content),
                    file_hash=file_hash,
                    mime_type=mime_type,
                    extracted_content=extracted_content,
                    summary=summary,
                    embedding=None,  # We use PGVector for embeddings now
                    category=category,
                    tags=tags,
                    user_id=user_id,
                    knowledge_base_id=knowledge_base_id
                )

                self.db.add(document)
                await self.db.commit()
                await self.db.refresh(document)

                # Create document chunks using PGVector
                if extracted_content:
                    await self._create_document_chunks_with_pgvector(document, extracted_content)

                logger.info(f"Document uploaded and processed successfully: {document.id}")
                return document

            finally:
                # Clean up temporary file
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)

        except Exception as e:
            logger.error(f"Error uploading document: {e}")
            await self.db.rollback()
            raise

    async def _create_document_chunks_with_pgvector(self, document: Document, content: str) -> None:
        """Create document chunks using PGVector for embeddings"""
        try:
            # text_chunks = self.get_semantic_split_points(content)
            text_chunks = await self._split_text(content)
            if not text_chunks:
                logger.warning(f"No text chunks provided for document {document.id}")
                return

            # Create LangChain documents with metadata
            langchain_docs = []
            for i, chunk_text in enumerate(text_chunks):
                doc = LangChainDocument(
                    page_content=chunk_text,
                    metadata={
                        "document_id": str(document.id),
                        "knowledge_base_id": str(document.knowledge_base_id),
                        "chunk_index": i,
                        "chunk_size": len(chunk_text),
                        "filename": document.filename,
                        "category": document.category or "general",
                        "file_path": document.file_path,
                        "mime_type": document.mime_type
                    }
                )
                langchain_docs.append(doc)

            # Get or create vector store
            collection_name = f"document_chunks_{document.user_id}".replace("-", "_")

            vector_store = PGVector(
                    connection=self.connection_string,
                    embeddings=self.embeddings,
                    collection_name=collection_name,
                    use_jsonb=True
                )


            if vector_store:
                collection_name = f"kb_{document.knowledge_base_id}" if document.knowledge_base_id else "default"
                logger.info(f"Adding {len(langchain_docs)} documents to PGVector collection: {collection_name}")
                
                # Add documents to vector store (synchronous operation)
                try:
                    vector_store.add_documents(langchain_docs)
                    logger.info(f"Successfully added {len(langchain_docs)} documents to PGVector collection: {collection_name}")
                except Exception as e:
                    logger.error(f"Error adding documents to PGVector: {e}")
                    raise
                
                # No longer creating DocumentChunk records - using langchain_pg_embedding directly
                await self.db.commit()
                logger.info(f"Created {len(text_chunks)} chunks for document {document.id} in PGVector")
                
        except Exception as e:
            logger.error(f"Error creating document chunks: {e}")
            await self.db.rollback()
            raise

    async def search_documents(
        self,
        query: str,
        user_id: UUID,
        limit: int = 10,
        knowledge_base_id: Optional[UUID] = None,
        category: Optional[str] = None,
        similarity_threshold: float = 0.8
    ) -> List[Dict[str, Any]]:
        """Search documents (wrapper for semantic_search)"""
        return await self.semantic_search(
            query=query,
            user_id=user_id,
            limit=limit,
            knowledge_base_id=knowledge_base_id,
            category=category,
            similarity_threshold=similarity_threshold
        )

    async def semantic_search(
        self,
        query: str,
        user_id: UUID,
        limit: int = 10,
        knowledge_base_id: Optional[UUID] = None,
        category: Optional[str] = None,
        similarity_threshold: float = 0.8
    ) -> List[Dict[str, Any]]:
        """Perform semantic search using PGVector"""
        try:
            collection_name = f"document_chunks_{user_id}".replace("-", "_")
            
            # Connect to vector store
            vector_store = PGVector(
                connection=self.connection_string,
                embeddings=self.embeddings,
                collection_name=collection_name,
                use_jsonb=True
            )
            
            # Build filter conditions
            filter_conditions = {}
            if knowledge_base_id:
                filter_conditions["knowledge_base_id"] = str(knowledge_base_id)
            if category:
                filter_conditions["category"] = category
            
            # Perform similarity search
            results = vector_store.similarity_search_with_score(
                query=query,
                k=limit,
                filter=filter_conditions if filter_conditions else None
            )
            
            # Format results
            search_results = []
            for doc, score in results:
                if score <= similarity_threshold:
                    search_results.append({
                        "document_id": doc.metadata.get("document_id"),
                        "filename": doc.metadata.get("filename"),
                        "content": doc.page_content,
                        "chunk_index": doc.metadata.get("chunk_index"),
                        "category": doc.metadata.get("category"),
                        "similarity": float(score),
                        "metadata": doc.metadata
                    })
            
            return search_results
            
        except Exception as e:
            logger.error(f"Error in semantic search: {e}")
            # Fallback to text search
            return await self._fallback_text_search(query, user_id, limit, knowledge_base_id, category)

    async def _fallback_text_search(
        self,
        query: str,
        user_id: UUID,
        limit: int,
        knowledge_base_id: Optional[UUID] = None,
        category: Optional[str] = None
    ) -> List[Dict[str, Any]]:
        """Fallback text search when vector search fails"""
        try:
            base_query = select(Document).where(Document.user_id == user_id)
            
            if knowledge_base_id:
                base_query = base_query.where(Document.knowledge_base_id == knowledge_base_id)
            
            if category:
                base_query = base_query.where(Document.category == category)
            
            # Simple text search
            text_query = base_query.where(
                Document.extracted_content.ilike(f"%{query}%")
            ).limit(limit)
            
            result = await self.db.execute(text_query)
            documents = result.scalars().all()
            
            # Format results
            search_results = []
            for doc in documents:
                search_results.append({
                    "document_id": str(doc.id),
                    "filename": doc.filename,
                    "content": doc.extracted_content[:500] if doc.extracted_content else "",
                    "summary": doc.summary,
                    "category": doc.category,
                    "similarity": 0.5,  # Default similarity for text search
                    "created_at": doc.created_at.isoformat()
                })
            
            return search_results
            
        except Exception as e:
            logger.error(f"Error in fallback text search: {e}")
            return []

    async def get_user_documents(
        self,
        user_id: UUID,
        skip: int = 0,
        limit: int = 100,
        category: Optional[str] = None,
        knowledge_base_id: Optional[UUID] = None
    ) -> List[Document]:
        """Get user documents with optional filtering"""
        try:
            query = select(Document).where(Document.user_id == user_id)
            
            if category:
                query = query.where(Document.category == category)
            
            if knowledge_base_id:
                query = query.where(Document.knowledge_base_id == knowledge_base_id)
            
            query = query.offset(skip).limit(limit).order_by(desc(Document.created_at))
            
            result = await self.db.execute(query)
            return result.scalars().all()
            
        except Exception as e:
            logger.error(f"Error getting user documents: {e}")
            raise

    async def get_by_id(self, document_id: str) -> Optional[Document]:
        """Get document by ID"""
        try:
            query = select(Document).where(Document.id == document_id)
            result = await self.db.execute(query)
            return result.scalar_one_or_none()
        except Exception as e:
            logger.error(f"Error getting document by ID: {e}")
            return None

    async def get_document_chunks(self, document_id: str, user_id: UUID) -> List[Dict[str, Any]]:
        """Get document chunks for preview from langchain_pg_embedding table"""
        try:
            # First check if document exists and user has permission
            document = await self.get_by_id(document_id)
            if not document:
                raise ValueError("Document not found")
            
            if document.user_id != user_id:
                raise ValueError("Permission denied")
            
            # Query langchain_pg_embedding table directly
            query = text("""
                SELECT 
                    id,
                    document,
                    cmetadata
                FROM langchain_pg_embedding 
                WHERE cmetadata->>'document_id' = :document_id
                ORDER BY CAST(cmetadata->>'chunk_index' AS INTEGER)
            """)
            
            result = await self.db.execute(query, {"document_id": str(document_id)})
            rows = result.fetchall()
            
            # Format chunks for response
            formatted_chunks = []
            for row in rows:
                metadata = row.cmetadata if row.cmetadata else {}
                formatted_chunks.append({
                    "id": str(row.id),
                    "content": row.document,
                    "chunk_index": metadata.get("chunk_index", 0),
                    "chunk_size": metadata.get("chunk_size", len(row.document) if row.document else 0),
                    "metadata": metadata
                })
            
            return formatted_chunks
            
        except Exception as e:
            logger.error(f"Error getting document chunks: {e}")
            raise

    async def delete(self, document: Document) -> None:
        """Delete a document and its chunks from langchain_pg_embedding"""
        try:
            # Delete document chunks from langchain_pg_embedding table
            delete_query = text("""
                DELETE FROM langchain_pg_embedding 
                WHERE cmetadata->>'document_id' = :document_id
            """)
            await self.db.execute(delete_query, {"document_id": str(document.id)})
            
            # Delete the document
            await self.db.delete(document)
            await self.db.commit()
            
            # Clean up file if it exists
            if document.file_path and os.path.exists(document.file_path):
                os.unlink(document.file_path)
                
            logger.info(f"Deleted document {document.id} and its embeddings")
            
        except Exception as e:
            logger.error(f"Error deleting document: {e}")
            await self.db.rollback()
            raise

    async def _get_document_by_hash(
        self,
        file_hash: str,
        user_id: UUID
    ) -> Optional[Document]:
        """Get document by file hash"""
        query = select(Document).where(
            Document.file_hash == file_hash,
            Document.user_id == user_id
        )
        result = await self.db.execute(query)
        return result.scalar_one_or_none()

    def _get_mime_type(self, filename: str) -> str:
        """Determine MIME type from filename"""
        extension = filename.lower().split('.')[-1] if '.' in filename else ''
        mime_types = {
            'pdf': 'application/pdf',
            'doc': 'application/msword',
            'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'txt': 'text/plain',
            'md': 'text/markdown',
            'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            'xls': 'application/vnd.ms-excel'
        }
        return mime_types.get(extension, 'application/octet-stream')

    async def _save_temp_file(self, content: bytes, filename: str) -> str:
        """Save file temporarily for processing"""
        temp_dir = tempfile.gettempdir()
        temp_file_path = os.path.join(temp_dir, f"temp_{filename}")
        
        with open(temp_file_path, 'wb') as f:
            f.write(content)
        
        return temp_file_path

    async def _save_file(self, content: bytes, filename: str, user_id: UUID) -> str:
        """Save file permanently"""
        # Create user-specific upload directory
        upload_dir = os.path.join(settings.UPLOAD_DIR, str(user_id))
        os.makedirs(upload_dir, exist_ok=True)
        
        # Generate unique filename
        file_path = os.path.join(upload_dir, filename)
        counter = 1
        base_name, extension = os.path.splitext(filename)
        
        while os.path.exists(file_path):
            new_filename = f"{base_name}_{counter}{extension}"
            file_path = os.path.join(upload_dir, new_filename)
            counter += 1
        
        # Save file
        with open(file_path, 'wb') as f:
            f.write(content)
        
        return file_path

    async def _extract_text_content(self, file_path: str, mime_type: str) -> str:
        """Extract text content from file"""
        try:
            if mime_type == 'text/plain' or mime_type == 'text/markdown':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            
            elif mime_type == 'application/pdf':
                text = ""
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
                return text.strip()
            
            elif mime_type in ['application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
                if mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                    doc = DocxDocument(file_path)
                    text = ""
                    for paragraph in doc.paragraphs:
                        text += paragraph.text + "\n"
                    return text.strip()
                else:
                    # For .doc files, we'd need python-docx2txt or similar
                    logger.warning(f"Unsupported document format: {mime_type}")
                    return ""
            
            else:
                logger.warning(f"Unsupported file type: {mime_type}")
                return ""
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return ""

    async def get_document_by_id(
        self,
        document_id: str,
        user_id: UUID
    ) -> Optional[Document]:
        """Get a document by ID for a specific user"""
        try:
            # Convert string ID to UUID
            doc_uuid = UUID(document_id)
            
            query = select(Document).where(
                Document.id == doc_uuid,
                Document.user_id == user_id
            )
            
            result = await self.db.execute(query)
            return result.scalar_one_or_none()
            
        except (ValueError, TypeError) as e:
            logger.error(f"Invalid document ID format {document_id}: {e}")
            return None
        except Exception as e:
            logger.error(f"Error getting document {document_id}: {e}")
            return None