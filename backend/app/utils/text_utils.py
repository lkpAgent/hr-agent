"""
Text utility functions
"""
import re
import html
from typing import List, Optional
import unicodedata
import logging

logger = logging.getLogger(__name__)


def clean_text(text: str) -> str:
    """Clean and normalize text"""
    if not text:
        return ""
    
    # Remove HTML tags
    text = remove_html_tags(text)
    
    # Normalize unicode
    text = unicodedata.normalize('NFKC', text)
    
    # Remove extra whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Strip leading/trailing whitespace
    text = text.strip()
    
    return text


def remove_html_tags(text: str) -> str:
    """Remove HTML tags from text"""
    if not text:
        return ""
    
    # Unescape HTML entities first
    text = html.unescape(text)
    
    # Remove HTML tags
    clean = re.compile('<.*?>')
    text = re.sub(clean, '', text)
    
    return text


def normalize_text(text: str) -> str:
    """Normalize text for search and comparison"""
    if not text:
        return ""
    
    # Convert to lowercase
    text = text.lower()
    
    # Remove punctuation and special characters
    text = re.sub(r'[^\w\s]', ' ', text)
    
    # Remove extra whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Strip
    text = text.strip()
    
    return text


def truncate_text(text: str, max_length: int, suffix: str = "...") -> str:
    """Truncate text to maximum length with suffix"""
    if not text or len(text) <= max_length:
        return text
    
    # Account for suffix length
    truncate_length = max_length - len(suffix)
    
    if truncate_length <= 0:
        return suffix[:max_length]
    
    # Try to truncate at word boundary
    truncated = text[:truncate_length]
    last_space = truncated.rfind(' ')
    
    if last_space > truncate_length * 0.8:  # If we can save 80% of the text
        truncated = truncated[:last_space]
    
    return truncated + suffix


def extract_keywords(text: str, max_keywords: int = 10) -> List[str]:
    """Extract keywords from text (simple implementation)"""
    if not text:
        return []
    
    # Normalize text
    normalized = normalize_text(text)
    
    # Split into words
    words = normalized.split()
    
    # Filter out common stop words (simple list)
    stop_words = {
        'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
        'of', 'with', 'by', 'is', 'are', 'was', 'were', 'be', 'been', 'have',
        'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could', 'should',
        'may', 'might', 'must', 'can', 'this', 'that', 'these', 'those',
        'i', 'you', 'he', 'she', 'it', 'we', 'they', 'me', 'him', 'her',
        'us', 'them', 'my', 'your', 'his', 'our', 'their'
    }
    
    # Filter words
    keywords = []
    for word in words:
        if (len(word) > 2 and 
            word not in stop_words and 
            word.isalpha()):
            keywords.append(word)
    
    # Count frequency and get most common
    word_freq = {}
    for word in keywords:
        word_freq[word] = word_freq.get(word, 0) + 1
    
    # Sort by frequency and return top keywords
    sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
    
    return [word for word, freq in sorted_words[:max_keywords]]


def split_text_into_chunks(
    text: str, 
    chunk_size: int = 1000, 
    overlap: int = 100
) -> List[str]:
    """Split text into overlapping chunks"""
    if not text:
        return []
    
    if len(text) <= chunk_size:
        return [text]
    
    chunks = []
    start = 0
    
    while start < len(text):
        end = start + chunk_size
        
        # If this is not the last chunk, try to end at a sentence boundary
        if end < len(text):
            # Look for sentence endings
            sentence_end = text.rfind('.', start, end)
            if sentence_end == -1:
                sentence_end = text.rfind('!', start, end)
            if sentence_end == -1:
                sentence_end = text.rfind('?', start, end)
            
            # If we found a sentence ending, use it
            if sentence_end > start + chunk_size * 0.5:
                end = sentence_end + 1
            else:
                # Otherwise, try to end at a word boundary
                space_pos = text.rfind(' ', start, end)
                if space_pos > start + chunk_size * 0.5:
                    end = space_pos
        
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        
        # Move start position with overlap
        start = end - overlap
        
        # Ensure we don't go backwards
        if start <= chunks[-1] if chunks else 0:
            start = end
    
    return chunks


def calculate_text_similarity(text1: str, text2: str) -> float:
    """Calculate simple text similarity based on common words"""
    if not text1 or not text2:
        return 0.0
    
    # Normalize texts
    words1 = set(normalize_text(text1).split())
    words2 = set(normalize_text(text2).split())
    
    if not words1 or not words2:
        return 0.0
    
    # Calculate Jaccard similarity
    intersection = len(words1.intersection(words2))
    union = len(words1.union(words2))
    
    return intersection / union if union > 0 else 0.0


def format_file_size(size_bytes: int) -> str:
    """Format file size in human readable format"""
    if size_bytes == 0:
        return "0 B"
    
    size_names = ["B", "KB", "MB", "GB", "TB"]
    i = 0
    size = float(size_bytes)
    
    while size >= 1024.0 and i < len(size_names) - 1:
        size /= 1024.0
        i += 1
    
    return f"{size:.1f} {size_names[i]}"


def escape_markdown(text: str) -> str:
    """Escape markdown special characters"""
    if not text:
        return ""
    
    # Characters that need escaping in markdown
    escape_chars = ['\\', '`', '*', '_', '{', '}', '[', ']', '(', ')', '#', '+', '-', '.', '!']
    
    for char in escape_chars:
        text = text.replace(char, f'\\{char}')
    
    return text


def generate_slug(text: str, max_length: int = 50) -> str:
    """Generate URL-friendly slug from text"""
    if not text:
        return ""
    
    # Convert to lowercase and normalize
    slug = text.lower()
    slug = unicodedata.normalize('NFKD', slug)
    
    # Remove non-ASCII characters
    slug = slug.encode('ascii', 'ignore').decode('ascii')
    
    # Replace spaces and special characters with hyphens
    slug = re.sub(r'[^a-z0-9]+', '-', slug)
    
    # Remove leading/trailing hyphens
    slug = slug.strip('-')
    
    # Truncate if too long
    if len(slug) > max_length:
        slug = slug[:max_length].rstrip('-')
    
    return slug or "untitled"


def mask_sensitive_data(text: str, mask_char: str = "*") -> str:
    """Mask sensitive data like emails, phone numbers, etc."""
    if not text:
        return ""
    
    # Mask email addresses
    text = re.sub(
        r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
        lambda m: m.group(0)[:2] + mask_char * (len(m.group(0)) - 4) + m.group(0)[-2:],
        text
    )
    
    # Mask phone numbers (simple pattern)
    text = re.sub(
        r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b',
        lambda m: mask_char * len(m.group(0)),
        text
    )
    
    return text


async def extract_text_content(file_path: str, mime_type: str) -> str:
    """Extract text content from a file path using robust handlers.

    This mirrors the enhanced document service's extraction logic so it can be reused
    across knowledge base ingestion and resume screening.
    """
    try:
        if mime_type in ('text/plain', 'text/markdown'):
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()

        elif mime_type == 'application/pdf':
            text = ""
            try:
                import PyPDF2  # type: ignore
            except Exception as e:
                logger.error(f"PyPDF2 not available for PDF extraction: {e}")
                return ""

            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    page_text = page.extract_text() or ""
                    if page_text:
                        text += page_text + "\n"
            return text.strip()

        elif mime_type in ('application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'):
            if mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                # Prefer docx2txt for robust extraction
                try:
                    import docx2txt  # type: ignore
                    extracted = docx2txt.process(file_path) or ""
                    if extracted.strip():
                        logger.info("Extracted DOCX content using docx2txt")
                        return extracted.strip()
                    else:
                        logger.info("docx2txt returned empty; falling back to python-docx")
                except Exception as e:
                    logger.info(f"docx2txt not available or failed ({e}); falling back to python-docx")

                # Fallback to python-docx
                try:
                    from docx import Document as DocxDocument  # type: ignore
                    doc = DocxDocument(file_path)
                    parts = []
                    for p in doc.paragraphs:
                        if p.text and p.text.strip():
                            parts.append(p.text)
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text and cell.text.strip():
                                    parts.append(cell.text)
                    return "\n".join(parts).strip()
                except Exception as e:
                    logger.error(f"python-docx failed to extract DOCX content: {e}")
                    return ""

            # .doc legacy format: unsupported here without external tools
            logger.warning(f"Unsupported .doc format for direct extraction: {mime_type}")
            return ""

        else:
            logger.warning(f"Unsupported file type: {mime_type}")
            return ""

    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {e}")
        return ""