"""
Unified document processing utilities for text extraction and MIME detection
"""
import os
import re
import tempfile
import logging
from typing import Optional

import fitz
import olefile
from docx import Document

logger = logging.getLogger(__name__)


def get_mime_type_from_filename(filename: str) -> str:
    ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''
    table = {
        'pdf': 'application/pdf',
        'doc': 'application/msword',
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'txt': 'text/plain',
        'md': 'text/markdown',
        'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        'xls': 'application/vnd.ms-excel'
    }
    return table.get(ext, 'application/octet-stream')


def normalize_text(text: str) -> str:
    if not text:
        return ''
    return ' '.join(text.replace('\r', '\n').split())

async def extract_text_with_advanced_formatting( html_content: str) -> str:
    """高级格式保留：区分段落、列表、标题等"""
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(html_content, 'html.parser')

    # 移除不需要的元素
    for element in soup(["script", "style", "head", "title", "meta"]):
        element.decompose()

    # 处理特定标签的格式
    formatting_rules = {
        'p': '\n\n',  # 段落：双换行
        'br': '\n',  # 换行：单换行
        'div': '\n',  # 分区：单换行
        'li': '\n• ',  # 列表项：换行+项目符号
        'h1': '\n\n# ',  # 一级标题
        'h2': '\n\n## ',  # 二级标题
        'h3': '\n\n### ',  # 三级标题
        'blockquote': '\n> '  # 引用块
    }

    # 应用格式规则
    for tag, prefix in formatting_rules.items():
        for element in soup.find_all(tag):
            # 在元素内容前添加格式前缀
            element.insert_before(prefix)

    # 获取文本
    text = soup.get_text()
    # 清理文本：合并多余空行，保留合理的段落间距
    lines = []
    empty_line_count = 0
    for line in text.splitlines():
        stripped_line = line.strip()
        if stripped_line:
            lines.append(stripped_line)
            empty_line_count = 0
        elif empty_line_count < 2:  # 最多保留两个连续空行
            lines.append('')
            empty_line_count += 1
    return '\n'.join(lines)
async def extract_text_from_file(file_path: str, mime_type: Optional[str] = None) -> str  :
    """Extract text content from a file path based on mime type.
    Supports: txt, md, pdf, docx, doc
    """
    try:
        print("开始解析文件:",file_path)
        mime = mime_type or get_mime_type_from_filename(os.path.basename(file_path))
        print("文件类型:",mime)
        if mime in ('text/plain', 'text/markdown'):
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            except UnicodeDecodeError:
                with open(file_path, 'rb') as f:
                    return f.read().decode('latin-1', errors='ignore')

        if mime == 'application/pdf':
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    text = ''
                    for page in pdf.pages:
                        page_text = page.extract_text() or ''
                        text += (page_text + '\n')
                return text.strip()
            except Exception as e:
                logger.error('pdfplumber extract pdf error,will use fitz')
                try:
                    text = ''
                    doc = fitz.open(file_path)
                    for page_num in range(len(doc)):
                        page = doc[page_num]
                        page_text = page.get_text()
                        if page_text and page_text.strip():
                            text += page_text + "\n"
                    doc.close()
                    if text.strip():
                        print("fitz解析成功")
                        return text.strip()
                    print(text)
                except Exception as e:
                    print(f"fitz抽取pdf失败: {e}")


        if mime == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':

            from docx import Document as DocxDocument
            try:
                doc = DocxDocument(file_path)
                parts = []
                for p in doc.paragraphs:
                    if p.text:
                        parts.append(p.text)
                if parts:
                    print("docx解析成功")
                    return '\n'.join(parts).strip()
                else:
                    import docx2txt
                    try:
                        text = docx2txt.process(file_path)
                        return text.strip() if text else ''
                    except Exception as e:
                        print(f"docx2txt 读取失败: {e}")
                        return ''
            except Exception as e:
                print(f"docx抽取失败: {e}")
                return ''

        if mime == 'application/msword':
            # .doc legacy, use docx2txt if available

            import docx2txt
            try:
                text = docx2txt.process(file_path)
                return (text or '').strip()
            except Exception as e:
                logger.warning(f'docx2txt failed: {e}')
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        html_content = f.read()
                        text = await extract_text_with_advanced_formatting(html_content=html_content)
                        return text
                except Exception as e:
                    logger.warning(f'Failed to extract text from {file_path}: {e}')
                    """提取.docx文件文本"""
                    try:
                        doc = Document(file_path)
                        text = ""
                        for paragraph in doc.paragraphs:
                            text += paragraph.text + "\n"
                        return text.strip()
                    except Exception as e:
                        print(f"DOCX提取失败: {e}")

                        text = extract_text_from_old_doc(file_path)
                        return text


        logger.warning(f'Unsupported mime type for extraction: {mime}')
        return ''
    except Exception as e:
        logger.error(f'Error extracting text from {file_path}: {e}')
        return ''


async def extract_text_from_bytes(file_bytes: bytes, filename: str) -> str:
    """Convenience helper to write bytes to a temp file and extract text"""
    temp_dir = tempfile.gettempdir()
    temp_path = os.path.join(temp_dir, f"_tmp_{filename}")
    try:
        with open(temp_path, 'wb') as f:
            f.write(file_bytes)
        return await extract_text_from_file(temp_path, get_mime_type_from_filename(filename))
    finally:
        try:
            if os.path.exists(temp_path):
                os.unlink(temp_path)
        except Exception:
            print(f"Failed to delete temp file: {temp_path}")
            pass

### 处理部分老式的doc文档

def extract_text_from_old_doc(file_path):
    """
    从 .doc 文件中提取文本内容，保持基本分段
    传入文件路径，返回文本内容字符串
    """
    try:
        # 检查文件是否为有效的 OLE 文档
        if not olefile.isOleFile(file_path):
            return ''

        ole = olefile.OleFileIO(file_path)
        text_content = ""

        # 检查是否存在 WordDocument 流
        if ole.exists('WordDocument'):
            stream = ole.openstream('WordDocument')
            data = stream.read()
            text_content = parse_word_document_content(data)

        ole.close()
        return text_content if text_content.strip() else ''

    except Exception as e:
        print(f"提取文本失败: {e}")
        return ''


def parse_word_document_content(data):
    """
    解析 WordDocument 流的二进制数据，提取文本内容
    """
    text_parts = []

    # 尝试不同的编码和偏移量来解码文本
    encoding_attempts = [
        ('utf-16-le', 0x400),  # 最常见的编码和偏移
        ('utf-16-le', 0x200),
        ('utf-16-le', 0x600),
        ('utf-8', 0x400),
        ('latin-1', 0x400),
    ]

    for encoding, offset in encoding_attempts:
        try:
            if offset < len(data):
                decoded_text = data[offset:].decode(encoding, errors='ignore')
                cleaned_text = clean_and_preserve_paragraphs(decoded_text)
                if is_meaningful_text(cleaned_text):
                    return cleaned_text
        except:
            continue

    # 如果上述方法都失败，尝试二进制扫描
    return extract_text_by_binary_scan(data)


def clean_and_preserve_paragraphs(text):
    """
    清理文本并保持基本的分段结构
    """
    # 替换各种换行符为统一的换行符
    text = re.sub(r'\r\n', '\n', text)  # Windows 换行
    text = re.sub(r'\r', '\n', text)  # Mac 换行

    # 移除控制字符和特殊字符，但保留换行符
    text = re.sub(r'[\x00-\x09\x0b-\x1f\x7f-\x9f]', ' ', text)

    # 合并过多的连续空白字符，但保留换行符
    lines = text.split('\n')
    cleaned_lines = []

    for line in lines:
        # 清理每行的空白
        line = re.sub(r'[^\S\n]+', ' ', line.strip())
        if line and is_meaningful_line(line):
            cleaned_lines.append(line)

    # 重新组合文本，保留段落结构
    result = str('\n'.join(cleaned_lines))

    # 确保文本以合理的分段结束
    return result.strip()


def is_meaningful_text(text):
    """
    判断文本是否有意义（包含足够的可读内容）
    """
    if len(text) < 50:  # 文本太短
        return False

    # 计算字母数字字符的比例
    alnum_count = sum(1 for c in text if c.isalnum())
    if alnum_count / len(text) < 0.1:  # 至少10%是字母数字
        return False

    # 检查是否包含常见单词
    common_words = ['的', '是', '在', '和', '有', '为', '了', '中', '与', '就',
                    'the', 'and', 'is', 'in', 'of', 'to', 'a', 'for', 'with', 'on']

    word_count = sum(1 for word in common_words if word in text.lower())
    return word_count >= 3  # 至少包含3个常见单词


def is_meaningful_line(line):
    """
    判断单行文本是否有意义
    """
    if len(line.strip()) < 2:  # 行太短
        return False

    # 计算字母数字字符比例
    alnum_count = sum(1 for c in line if c.isalnum())
    return alnum_count / len(line) > 0.2  # 至少20%是字母数字


def extract_text_by_binary_scan(data):
    """
    通过二进制扫描提取文本（备用方法）
    """
    text_parts = []
    current_word = []

    for byte in data:
        # ASCII 可打印字符
        if 32 <= byte <= 126:
            current_word.append(chr(byte))
        else:
            # 遇到非打印字符时，保存当前单词
            if len(current_word) >= 3:  # 至少3个字符才认为是单词
                word = ''.join(current_word)
                if any(c.isalpha() for c in word):  # 包含至少一个字母
                    text_parts.append(word)
            current_word = []

    # 处理最后一个单词
    if len(current_word) >= 3:
        word = ''.join(current_word)
        if any(c.isalpha() for c in word):
            text_parts.append(word)

    # 将单词组合成文本，每10个单词换行
    if text_parts:
        lines = []
        for i in range(0, len(text_parts), 10):
            lines.append(' '.join(text_parts[i:i + 10]))
        return '\n'.join(lines)

    return ""


### 老式word文档处理结束